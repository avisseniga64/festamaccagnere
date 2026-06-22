import xlrd
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

script_dir = os.path.dirname(os.path.abspath(__file__))
xls_path = os.path.join(script_dir, "MENU FESTA 2026_STAMPA.xls")

# Leggi dati dal foglio MENU BAR
wb = xlrd.open_workbook(xls_path)
sheet = wb.sheet_by_name('MENU BAR')

# Raccogli dati raggruppati per sezione (righe vuote = separatore)
sections = []
current_items = []

for row_idx in range(sheet.nrows):
    col0 = str(sheet.cell(row_idx, 0).value).strip()
    col1_val = sheet.cell(row_idx, 1).value

    if col0 == "":
        if current_items:
            sections.append(current_items)
            current_items = []
        continue

    if col1_val:
        price = float(col1_val)
        price_str = f"\u20ac {price:.2f}".replace(".", ",")
        display_name = col0.title()
        # Fix preposizioni
        for word in ['Di', 'Al', 'Alla', 'Alle', 'Del', 'Delle', 'Con', 'E', 'In', 'Da']:
            display_name = display_name.replace(f' {word} ', f' {word.lower()} ')
        current_items.append((display_name, price_str))

if current_items:
    sections.append(current_items)

total_items = sum(len(s) for s in sections)
print(f"Sezioni: {len(sections)}, Voci totali: {total_items}")

# --- Crea documento Word A4 portrait ---
doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.orientation = WD_ORIENT.PORTRAIT
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# Larghezza utile: 210 - 25 - 25 = 160mm
TAB_STOP_POS = Mm(160)

color_rosso = RGBColor(0x8B, 0x1A, 0x1A)
color_marrone = RGBColor(0x6B, 0x44, 0x23)
color_grigio = RGBColor(0x8B, 0x6B, 0x4A)

# Footer
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("In cassa \u00e8 disponibile il libro ingredienti con i componenti e gli allergeni di ogni singola preparazione alimentare. In caso di bisogno rivolgersi al personale.")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = color_grigio
run.font.name = 'Georgia'

# Titolo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(0)
p.space_after = Pt(4)
run = p.add_run("MENU BAR")
run.font.size = Pt(32)
run.font.color.rgb = color_rosso
run.font.bold = True
run.font.name = 'Georgia'

# Sottotitolo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(0)
p.space_after = Pt(2)
run = p.add_run("Festa delle Maccagnere")
run.font.size = Pt(14)
run.font.color.rgb = color_marrone
run.font.italic = True
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(0)
p.space_after = Pt(4)
run = p.add_run("2026")
run.font.size = Pt(18)
run.font.color.rgb = color_rosso
run.font.bold = True
run.font.name = 'Georgia'

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(4)
p.space_after = Pt(16)
run = p.add_run("\u2666 \u2022 \u2666 \u2022 \u2666")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xC4, 0x95, 0x6A)
run.font.name = 'Georgia'

# Voci con tab leader puntinato, separatore tra sezioni
for sec_idx, items in enumerate(sections):
    for name, price in items:
        p = doc.add_paragraph()
        p.space_before = Pt(3)
        p.space_after = Pt(3)

        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        run = p.add_run(name)
        run.font.size = Pt(12)
        run.font.name = 'Georgia'
        run.font.color.rgb = RGBColor(0x2C, 0x18, 0x10)

        p.add_run("\t")
        run_price = p.add_run(price)
        run_price.font.size = Pt(12)
        run_price.font.bold = True
        run_price.font.name = 'Georgia'
        run_price.font.color.rgb = color_rosso

    # Separatore tra sezioni (tranne dopo l'ultima)
    if sec_idx < len(sections) - 1:
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        p.space_after = Pt(6)

# Salva
output_path = os.path.join(script_dir, "Menu_Bar_2026.docx")
doc.save(output_path)
print(f"\nFile generato: {output_path}")
