import xlrd
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

script_dir = os.path.dirname(os.path.abspath(__file__))
xls_path = os.path.join(script_dir, "MENU FESTA 2026_STAMPA.xls")

# Leggi dati dal foglio XLS
wb = xlrd.open_workbook(xls_path)
sheet = None
for name in wb.sheet_names():
    if "listino grande" in name.lower():
        sheet = wb.sheet_by_name(name)
        break

# Raccogli dati
menu_items = []
bevande_items = []
nota_surgelato = ""
nota_footer = ""

current_section = None
for row_idx in range(sheet.nrows):
    col0 = str(sheet.cell(row_idx, 0).value).strip()
    col1_val = sheet.cell(row_idx, 1).value

    if col0 == "LISTINO PREZZI" or col0 == "":
        continue
    if col0 == "MENU'" and str(col1_val).strip().upper() == "EURO":
        current_section = "menu"
        continue
    elif col0 == "BEVANDE" and str(col1_val).strip().upper() == "EURO":
        current_section = "bevande"
        continue
    if "* PRODOTTO SURGELATO" in col0:
        nota_surgelato = "* Surgelato"
        continue
    if "In cassa" in col0:
        nota_footer = col0
        continue

    if col0 and col1_val:
        price = float(col1_val)
        price_str = f"\u20ac{price:.2f}".replace(".", ",")

        display_name = col0
        if "(SOLO SECONDO WEEKEND)" in col0.upper():
            display_name = col0.upper().replace("(SOLO SECONDO WEEKEND)", "").strip()
            display_name = display_name.title() + " (2\u00b0we)"
        elif "(SOLO PRIMO WEEKEND)" in col0.upper():
            display_name = col0.upper().replace("(SOLO PRIMO WEEKEND)", "").strip()
            display_name = display_name.title() + " (1\u00b0we)"
        else:
            display_name = col0.title()

        for word in ['Di', 'Al', 'Alla', 'Alle', 'Del', 'Delle', 'Con', 'E', 'In', 'Da']:
            display_name = display_name.replace(f' {word} ', f' {word.lower()} ')

        if current_section == "menu":
            menu_items.append((display_name, price_str))
        elif current_section == "bevande":
            bevande_items.append((display_name, price_str))

print(f"Menu: {len(menu_items)} voci")
print(f"Bevande: {len(bevande_items)} voci")

# --- Crea documento Word A5 LANDSCAPE ---
doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(148)
section.orientation = WD_ORIENT.LANDSCAPE
section.top_margin = Mm(3)
section.bottom_margin = Mm(3)
section.left_margin = Mm(4)
section.right_margin = Mm(4)

# Colori
color_rosso = RGBColor(0x8B, 0x1A, 0x1A)
color_marrone = RGBColor(0x6B, 0x44, 0x23)
color_grigio = RGBColor(0x8B, 0x6B, 0x4A)

# Nessun footer di sezione per risparmiare spazio
# Aggiungo piè di pagina Word
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run(nota_footer.replace('\n', ' '))
run.font.size = Pt(6)
run.font.italic = True
run.font.color.rgb = color_grigio
run.font.name = 'Arial Narrow'

# --- Titolo compatto su una riga ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(0)
p.space_after = Pt(2)
pf = p.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
run = p.add_run("FESTA DELLE MACCAGNERE 2026")
run.font.size = Pt(14)
run.font.color.rgb = color_rosso
run.font.bold = True
run.font.name = 'Arial Narrow'

# Costruisco una SINGOLA tabella con tutte le righe
# Costruisco una SINGOLA tabella con tutte le righe
# Colonna sinistra = menu, colonna destra = bevande
# Uso max(len(menu), len(bevande)) righe + 1 riga header (surgelato è nell'header)
max_rows = max(len(menu_items), len(bevande_items))
# +1 per intestazione
total_rows = max_rows + 1

table = doc.add_table(rows=total_rows, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Larghezze: Qty(8) | Nome(82) | Prezzo(15) || Qty(8) | Nome(68) | Prezzo(15)
col_widths = [Mm(7), Mm(80), Mm(14), Mm(7), Mm(72), Mm(14)]

def set_cell(cell, text, size=9, bold=False, color=None, align=None, font_name='Arial Narrow'):
    p = cell.paragraphs[0]
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = Pt(size + 3)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

# Riga 0: intestazioni con surgelato a fianco di MENÙ
set_cell(table.rows[0].cells[0], "Qt", 8, bold=True, color=color_grigio, align=WD_ALIGN_PARAGRAPH.CENTER)
# MENÙ + surgelato nella stessa cella
p = table.rows[0].cells[1].paragraphs[0]
p.space_before = Pt(0)
p.space_after = Pt(0)
pf = p.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.line_spacing = Pt(13)
run = p.add_run("MEN\u00d9  ")
run.font.size = Pt(10)
run.font.bold = True
run.font.name = 'Arial Narrow'
run.font.color.rgb = color_rosso
run2 = p.add_run(nota_surgelato)
run2.font.size = Pt(7)
run2.font.italic = True
run2.font.name = 'Arial Narrow'
run2.font.color.rgb = color_grigio

set_cell(table.rows[0].cells[2], "", 8)
set_cell(table.rows[0].cells[3], "Qt", 8, bold=True, color=color_grigio, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(table.rows[0].cells[4], "BEVANDE", 10, bold=True, color=color_rosso)
set_cell(table.rows[0].cells[5], "", 8)

# Righe dati
for i in range(max_rows):
    row = table.rows[i + 1]
    
    # Colonna sinistra (menu)
    if i < len(menu_items):
        name, price = menu_items[i]
        set_cell(row.cells[0], "___", 7, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[1], name, 9)
        set_cell(row.cells[2], price, 9, bold=True, color=color_rosso, align=WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Colonna destra (bevande)
    if i < len(bevande_items):
        name, price = bevande_items[i]
        set_cell(row.cells[3], "___", 7, color=RGBColor(0xCC, 0xCC, 0xCC), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[4], name, 9)
        set_cell(row.cells[5], price, 9, bold=True, color=color_rosso, align=WD_ALIGN_PARAGRAPH.RIGHT)

# (surgelato è già nell'intestazione)

# Imposta larghezze e stile tabella + altezza righe fissa
for row in table.rows:
    # Forza altezza riga esatta
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="295" w:hRule="exact"/>')
    trPr.append(trHeight)
    
    for idx, cell in enumerate(row.cells):
        cell.width = col_widths[idx]
        # Rimuovi margini celle
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            '<w:top w:w="0" w:type="dxa"/>'
            '<w:bottom w:w="0" w:type="dxa"/>'
            '<w:left w:w="14" w:type="dxa"/>'
            '<w:right w:w="14" w:type="dxa"/>'
            '</w:tcMar>'
        )
        tcPr.append(tcMar)

# Riga header leggermente più alta
tr0 = table.rows[0]._tr
trPr0 = tr0.get_or_add_trPr()
# Rimuovi vecchio trHeight se presente
for child in trPr0:
    if child.tag.endswith('trHeight'):
        trPr0.remove(child)
trHeight0 = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="350" w:hRule="exact"/>')
trPr0.append(trHeight0)

# Bordi tabella
tbl = table._tbl
tblPr = tbl.tblPr
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="8B1A1A"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="8B1A1A"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="dotted" w:sz="2" w:space="0" w:color="E0D5C8"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

# Aggiungi bordo verticale tra colonna 2 e 3 (separatore menu/bevande)
# Lo facciamo impostando bordo destro sulla cella 2 di ogni riga
for row in table.rows:
    tc = row.cells[2]._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D4A574"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

# Footer è già nel piè di pagina Word

# Salva
output_path = os.path.join(script_dir, "Menu_Ordine_A5.docx")
doc.save(output_path)
print(f"\nFile generato: {output_path}")
