import xlrd
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

script_dir = os.path.dirname(os.path.abspath(__file__))
xls_path = os.path.join(script_dir, "MENU FESTA 2026_STAMPA.xls")

# Leggi dati dal foglio XLS
wb = xlrd.open_workbook(xls_path)
sheet = None
for name in wb.sheet_names():
    if "listino grande" in name.lower():
        sheet = wb.sheet_by_name(name)
        break

# Raccogli dati
menu_items = []
bevande_items = []
nota_surgelato = ""
nota_footer = ""

current_section = None
for row_idx in range(sheet.nrows):
    col0 = str(sheet.cell(row_idx, 0).value).strip()
    col1_val = sheet.cell(row_idx, 1).value

    if col0 == "LISTINO PREZZI" or col0 == "":
        continue
    if col0 == "MENU'" and str(col1_val).strip().upper() == "EURO":
        current_section = "menu"
        continue
    elif col0 == "BEVANDE" and str(col1_val).strip().upper() == "EURO":
        current_section = "bevande"
        continue
    if "* PRODOTTO SURGELATO" in col0:
        nota_surgelato = "* Prodotto surgelato all'origine"
        continue
    if "In cassa" in col0:
        nota_footer = col0
        continue

    if col0 and col1_val:
        price = float(col1_val)
        price_str = f"\u20ac {price:.2f}".replace(".", ",")

        display_name = col0
        note = ""
        if "(SOLO SECONDO WEEKEND)" in col0.upper():
            display_name = col0.upper().replace("(SOLO SECONDO WEEKEND)", "").strip()
            display_name = display_name.title()
            note = " (solo secondo weekend)"
        elif "(SOLO PRIMO WEEKEND)" in col0.upper():
            display_name = col0.upper().replace("(SOLO PRIMO WEEKEND)", "").strip()
            display_name = display_name.title()
            note = " (solo primo weekend)"
        else:
            display_name = col0.title()

        for word in ['Di', 'Al', 'Alla', 'Alle', 'Del', 'Delle', 'Con', 'E', 'In', 'Da']:
            display_name = display_name.replace(f' {word} ', f' {word.lower()} ')

        if current_section == "menu":
            menu_items.append((display_name + note, price_str))
        elif current_section == "bevande":
            bevande_items.append((display_name + note, price_str))

print(f"Menu: {len(menu_items)} voci")
print(f"Bevande: {len(bevande_items)} voci")

# --- Crea documento Word A3 portrait ---
doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.orientation = WD_ORIENT.PORTRAIT
section.top_margin = Mm(15)
section.bottom_margin = Mm(15)
section.left_margin = Mm(20)
section.right_margin = Mm(20)

# Larghezza utile per tab stop: 210 - 20 - 20 = 170mm
TAB_STOP_POS = Mm(170)

color_rosso = RGBColor(0x8B, 0x1A, 0x1A)
color_marrone = RGBColor(0x6B, 0x44, 0x23)
color_grigio = RGBColor(0x8B, 0x6B, 0x4A)

# Footer nel piè di pagina
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run(nota_footer.replace('\n', ' '))
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = color_grigio
run.font.name = 'Georgia'

def add_title(doc, text, size=28):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color_rosso
    run.font.bold = True
    run.font.name = 'Georgia'

def add_subtitle(doc, text, size=14, italic=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color_marrone
    run.font.italic = italic
    run.font.name = 'Georgia'

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4)
    p.space_after = Pt(12)
    run = p.add_run("\u2666 \u2022 \u2666 \u2022 \u2666")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xC4, 0x95, 0x6A)
    run.font.name = 'Georgia'

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.color.rgb = color_rosso
    run.font.bold = True
    run.font.name = 'Georgia'
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="2" w:color="8B1A1A"/></w:pBdr>')
    pPr.append(pBdr)

def add_menu_items(doc, items):
    """Aggiunge voci menu usando tab con leader puntinato per allineare prezzi a destra"""
    for name, price in items:
        p = doc.add_paragraph()
        p.space_before = Pt(4)
        p.space_after = Pt(4)

        # Aggiungi tab stop allineato a destra con leader puntinato
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        # Nome piatto
        run = p.add_run(name)
        run.font.size = Pt(11)
        run.font.name = 'Georgia'
        run.font.color.rgb = RGBColor(0x2C, 0x18, 0x10)

        # Tab + prezzo
        run_tab = p.add_run("\t")
        run_price = p.add_run(price)
        run_price.font.size = Pt(11)
        run_price.font.bold = True
        run_price.font.name = 'Georgia'
        run_price.font.color.rgb = color_rosso

# ===== PAGINA 1: MENU =====
add_title(doc, "LISTINO PREZZI")
add_subtitle(doc, "Festa delle Maccagnere")
add_subtitle(doc, "2026", size=18, italic=False)
add_divider(doc)

add_section_title(doc, "Men\u00f9")
add_menu_items(doc, menu_items)

# Nota surgelato
p = doc.add_paragraph()
p.space_before = Pt(12)
run = p.add_run(nota_surgelato)
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = color_grigio
run.font.name = 'Georgia'

# ===== PAGE BREAK =====
doc.add_page_break()

# ===== PAGINA 2: BEVANDE =====
add_title(doc, "BEVANDE")
add_subtitle(doc, "Festa delle Maccagnere")
add_subtitle(doc, "2026", size=18, italic=False)
add_divider(doc)

add_section_title(doc, "Bevande")
add_menu_items(doc, bevande_items)

# Salva
output_path = os.path.join(script_dir, "Menu_Festa_2026_A3.docx")
doc.save(output_path)
print(f"\nFile generato: {output_path}")
